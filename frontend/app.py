# app.py
import streamlit as st
import uuid
import json
import sys
import os
import requests
import google.generativeai

from cookies import ensure_ready
cookies = ensure_ready()


# Add root dir so utils/ can be imported
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from ui import apply_custom_styles, show_header, logo_animation

from utils.firebase_config import db
from firebase_admin import firestore as fa_firestore

from streamlit.components.v1 import html as components_html
from utils.llm import generate_bot_config_gemini, chat_with_gemini

BACKEND="https://mentesav7.onrender.com"

from utils.file_handle import safe_text
# --- PAGE CONFIG ---
st.set_page_config(
    page_title="Mentesa",
    page_icon="frontend/logo.png",
    layout="centered",
    initial_sidebar_state="expanded"
)

st.markdown("""
    <meta name="description" content="Mentesa – A no-code platform to create, manage, and chat with your own AI bots.">
    <meta name="keywords" content="AI bots, chatbot builder, Mentesa, generative AI, Streamlit">
""", unsafe_allow_html=True)

from auth import auth_ui, require_login

st.set_page_config(page_title="Mentesa")


# --- Authentication check ---
if "user" not in st.session_state or not st.session_state["user"]:
    auth_ui()
    st.stop()

user = st.session_state["user"]
# For any page or operation that must be protected:
# user = require_login()
# uid = user['uid']
# now query Firestore for bots owned by uid, e.g.:
# db.collection("bots").where("owner_uid", "==", uid).stream()

# Logo Mentesa animation
logo_animation()

# Apply styles
apply_custom_styles()

# # Show header
#show_header()

# ---------------- utility: load bots for current user ----------------
def load_user_bots():
    """
    Return list of bot dicts owned by the current user.
    This will attempt to query by owner_uid and owner_email and deduplicate results.
    Always returns a list (possibly empty).
    """
    user = st.session_state.get("user") or {}
    uid = user.get("uid")
    email = user.get("email")
    bots = {}
    try:
        # Try owner_uid first (fast and strict)
        if uid:
            try:
                for doc in db.collection("bots").where("owner_uid", "==", uid).stream():
                    d = doc.to_dict() or {}
                    d["id"] = doc.id
                    bots[d["id"]] = d
            except Exception:
                # don't fail hard — continue to owner_email fallback
                pass

        # Then try owner_email (covers records created earlier or backend that stored only email)
        if email:
            try:
                for doc in db.collection("bots").where("owner_email", "==", email).stream():
                    d = doc.to_dict() or {}
                    d["id"] = doc.id
                    bots[d["id"]] = d
            except Exception:
                # fallback to filtering client-side below
                pass

        # If both queries yielded nothing, as a robust fallback we fetch a small set and filter locally.
        if not bots:
            try:
                # only fetch a limited set to avoid scanning entire DB; if your dataset is small you can remove limit
                for doc in db.collection("bots").stream():
                    d = doc.to_dict() or {}
                    owner_uid = d.get("owner_uid")
                    owner_email = d.get("owner_email")
                    if (uid and owner_uid == uid) or (email and owner_email == email):
                        d["id"] = doc.id
                        bots[d["id"]] = d
            except Exception:
                # last resort: return empty list
                pass
    except Exception as e:
        st.error(f"Failed to load your bots: {e}")
        return []

    # return deduplicated bots as list
    return list(bots.values())

# ---------------- BOT CREATION ----------------
from utils.llm import generate_bot_config_gemini

def create_and_save_bot():
    st.subheader("✨ Create Your Bot")
    st.write("Describe the bot you want, and we'll generate it with AI.")

    name = "Mentesa_Bot"
    prompt = st.text_area("🤔 What type of bot do you want?")
    url = st.text_input("🌐 (Optional) Website URL for the bot to ingest (include https://)")
    uploaded_files = st.file_uploader(
        "📂 (Optional) Upload RAG Files (PDF, DOCX, TXT)", 
        type=["pdf", "docx", "txt"], 
        accept_multiple_files=True
    )

    if st.button("🚀 Create Bot"):
        if not prompt.strip():
            st.warning("Please enter a prompt before generating.")
            return

        import time
        bot_name = name.strip() if name.strip() else f"Bot_{int(time.time())}"

        # Prepare files payload if files were uploaded
        files_payload = []
        if uploaded_files:
            import io
            for f in uploaded_files:
                try:
                    data_bytes = f.read()
                    filename = f.name
                    content = ""

                    ext = filename.lower().rsplit(".", 1)[-1]

                    if ext == "pdf":
                        from PyPDF2 import PdfReader
                        reader = PdfReader(io.BytesIO(data_bytes))
                        content = "\n".join([p.extract_text() or "" for p in reader.pages])
                    elif ext == "docx":
                        from docx import Document
                        doc = Document(io.BytesIO(data_bytes))
                        paragraphs = [p.text for p in doc.paragraphs]
                        content = "\n".join(paragraphs)
                    else:  # txt or fallback
                        try:
                            content = data_bytes.decode("utf-8")
                        except Exception:
                            content = data_bytes.decode("latin-1", errors="ignore")

                    from utils.file_handle import safe_text
                    content = safe_text(content)[:15000]
                    if not content.strip():
                        content = "-"

                    files_payload.append({"name": filename, "text": content})

                except Exception as e:
                    st.error(f"Failed to read uploaded file '{f.name}': {e}")
                    return

        # owner info
        user = st.session_state.get("user") or {}
        owner_uid = user.get("uid")
        owner_email = user.get("email")
        
        payload = {
            "name": bot_name.strip(),
            "prompt": prompt.strip(),
            "config": {"urls": [url.strip()]} if url.strip() else {},
            "files": files_payload,
            # --- owner info for backend/firestore
            "owner_uid": owner_uid,
            "owner_email": owner_email,
        }
        
        with st.spinner("Generating bot..."):
            try:
                response = requests.post(f"{BACKEND}/bots", json=payload, timeout=120)
            except Exception as e:
                st.error(f"Request failed: {e}")
                return

        if response.status_code in (200, 201):
            data = response.json()
            bot_name = data["bot"]["name"]
            if files_payload:
                uploaded_names = ", ".join([f['name'] for f in files_payload])
                st.success(f"✅ Bot '{bot_name}' created and saved with file(s): {uploaded_names}")
            else:
                st.success(f"✅ Bot '{bot_name}' created and saved!")
        else:
            st.error(f"Failed to save bot: {response.text}")

# ---------------- CHAT INTERFACE ----------------
def normalize_history(raw_history):
    """Convert old {'user':'', 'bot':''} into [{'role','content'}, ...]."""
    normalized = []
    for turn in raw_history:
        if isinstance(turn, dict):
            if "role" in turn and "content" in turn:
                normalized.append({"role": turn["role"], "content": turn["content"]})
            elif "user" in turn and "bot" in turn:
                normalized.append({"role": "user", "content": turn["user"]})
                normalized.append({"role": "bot", "content": turn["bot"]})
    return normalized

def chat_interface():
    st.header("💬 Chat with Your Bot")
    st.markdown("---")

    # Ensure we have user uid/email
    user = st.session_state.get("user") or {}
    uid = user.get("uid")
    email = user.get("email")
    if not (uid or email):
        st.error("User not found. Please sign in.")
        return

    # --- Load bots from Firebase (only current user's bots) ---
    try:
        bots = load_user_bots()
    except Exception as e:
        st.error(f"Failed to load bots: {e}")
        return

    if not bots:
        st.info("You don't have any bots yet — create one first.")
        return

    # --- Select bot ---
    selected_bot_info = st.selectbox(
        "Choose a bot",
        options=bots,
        format_func=lambda b: f"{b.get('name','(unknown)')} ({b.get('id','')[:6]})",
        key="chat_selectbox"
    )

    if not selected_bot_info:
        st.info("No bot selected.")
        return

    selected_bot_id = selected_bot_info["id"]

    st.markdown("---")

    # --- Load chat history from Firebase ---
    if "chat_bot_id" not in st.session_state or st.session_state.chat_bot_id != selected_bot_id:
        st.session_state.chat_bot_id = selected_bot_id
        history_doc = db.collection("bot_chats").document(selected_bot_id).get()
        if history_doc.exists:
            st.session_state.history = history_doc.to_dict().get("history", [])
        else:
            st.session_state.history = []

    history = st.session_state.history

    if "typing" not in st.session_state:
        st.session_state.typing = False

    # --- Iframe for chat UI ---
    history_json = json.dumps(history)
    typing_json = json.dumps(bool(st.session_state.typing))
    iframe_html = f"""
    <!doctype html>
    <html>
    <head>
      <meta charset="utf-8" />
      <style>
        body {{ margin:0; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial; }}
        .chat-box {{
            height: 500px;
            overflow-y: auto;
            padding: 12px;
            width: 100%;
            box-sizing: border-box;
            background: transparent;
        }}
        .msg-user {{
            background: #0084ff;
            color: white;
            padding: 8px 12px;
            border-radius: 12px;
            max-width: max-content;
            margin-bottom: 8px;
            margin-left: auto;
            text-align: right;
            white-space: pre-wrap;
            word-wrap: break-word;
        }}
        .msg-bot {{
            background: #f1f0f0;
            color: #000;
            padding: 8px 12px;
            border-radius: 12px;
            max-width: 70%;
            margin-bottom: 8px;
            margin-right: auto;
            text-align: left;
            white-space: pre-wrap;
            word-wrap: break-word;
        }}
        .msg-typing {{
            background: #f1f0f0;
            color: black;
            padding: 8px 12px;
            border-radius: 12px;
            max-width: 30%;
            margin-bottom: 8px;
            margin-right: auto;
            font-style: italic;
            opacity: 1;
        }}
      </style>
    </head>
    <body>
      <div id="chat-scroll-box" class="chat-box"></div>
      <script>
        const history = {history_json};
        const typing = {typing_json};
        function renderChat() {{
            const box = document.getElementById('chat-scroll-box');
            box.innerHTML = "";
            for (let i=0;i<history.length;i++) {{
                const turn = history[i];
                const div = document.createElement('div');
                div.className = turn.role === 'user' ? 'msg-user' : 'msg-bot';
                div.textContent = (turn.role === 'user' ? '🧑 ' : '🤖 ') + turn.content;
                box.appendChild(div);
            }}
            if (typing) {{
                const t = document.createElement('div');
                t.className = 'msg-typing';
                t.textContent = '🤖 typing…';
                box.appendChild(t);
            }}
            box.scrollTop = box.scrollHeight;
            setTimeout(()=>{{ box.scrollTop = box.scrollHeight; }}, 50);
            setTimeout(()=>{{ box.scrollTop = box.scrollHeight; }}, 300);
        }}
        setTimeout(renderChat, 10);
      </script>
    </body>
    </html>
    """
    components_html(iframe_html, height=460, width=1200, scrolling=True)

    # --- User input ---
    user_input = st.chat_input("Type your message…")
    if user_input:
        history.append({"role": "user", "content": user_input})
        st.session_state.history = history
        st.session_state.typing = True
        # Save immediately to Firebase
        db.collection("bot_chats").document(selected_bot_id).set({"history": history})
        st.rerun()

    # --- Generate bot reply ---
    if st.session_state.typing:
        response = requests.post(
            f"{BACKEND}/chat",
            json={"bot_id": selected_bot_id, "message": history[-1]["content"]},
            timeout=60
        )
        
        if response.status_code != 200:
            st.error(f"❌ Chat request failed: {response.status_code}\n{response.text}")
            return
        
        try:
            data = response.json()
            reply = data.get("reply", "⚠️ No reply received")
        except Exception as e:
            st.error(f"❌ Failed to parse JSON: {e}\nRaw response:\n{response.text}")
            return
        
        
        history.append({"role": "bot", "content": reply})
        st.session_state.history = history
        st.session_state.typing = False
        # Save reply to Firebase
        db.collection("bot_chats").document(selected_bot_id).set({"history": history})
        st.rerun()

# ---------------- BOT MANAGEMENT ----------------
def bot_management_ui():
    st.subheader("🛠️ Manage Your Bots")

    # Ensure we have user uid/email
    user = st.session_state.get("user") or {}
    uid = user.get("uid")
    email = user.get("email")
    if not (uid or email):
        st.error("User not found. Please sign in.")
        return

    # Load bots from Firebase (only this user's bots)
    try:
        bots = load_user_bots()
    except Exception as e:
        st.error(f"Failed to load bots: {e}")
        return

    if not bots:
        st.info("No bots available — create one first.")
        return

    # --- Select bot ---
    selected_bot_info = st.selectbox(
        "Choose a bot",
        options=bots,
        format_func=lambda b: f"{b.get('name','(unknown)')} ({b.get('id','')[:6]})",
        key="manage_select"
    )
    if not selected_bot_info:
        st.info("No bot selected.")
        return
    selected_bot_id = selected_bot_info["id"]

    # --- Bot management columns ---
    col1, col2, col3, col4 = st.columns([2, 3, 1, 1])

    # Rename
    new_name = col1.text_input("Name", value=selected_bot_info.get('name',''), key=f"name_{selected_bot_id}")
    if col1.button("✏️ Rename", key=f"rename_{selected_bot_id}"):
        db.collection("bots").document(selected_bot_id).update({"name": new_name})
        st.success("Renamed!")
        st.rerun()

    # Update Personality
    new_persona = col2.text_area("Personality", value=selected_bot_info.get('personality', ""), key=f"persona_{selected_bot_id}", height=80)
    if col2.button("✏️ Update", key=f"update_{selected_bot_id}"):
        db.collection("bots").document(selected_bot_id).update({"personality": new_persona})
        st.success("Personality updated!")
        st.rerun()

    # Clear Chat
    if col3.button("🧹 Clear Chat", key=f"manage_clear_{selected_bot_id}"):
        db.collection("bot_chats").document(selected_bot_id).delete()
        if "history" in st.session_state:
            st.session_state.history = []
        st.success("Chat history cleared!")
        st.rerun()

    # Delete Bot
    if col4.button("🗑️ Delete", key=f"delete_{selected_bot_id}"):
        db.collection("bots").document(selected_bot_id).delete()
        st.success("Bot deleted!")
        st.rerun()

    # --- RAG Management (Files & URLs) ---
    st.markdown("---")
    st.subheader("📂 Upload RAG Files")

    import io
    from utils.file_handle import upload_file, safe_text   # ensure safe_text is importable

    # Upload Files
    uploaded_file = st.file_uploader("Upload a RAG File", key=f"file_{selected_bot_id}")
    if uploaded_file:
        filename = uploaded_file.name
        flag = f"uploaded_{selected_bot_id}_{filename}"
        success_flag = f"success_{selected_bot_id}"

        # If file already processed in this session
        if st.session_state.get(flag):
            # Show success if it was freshly uploaded
            if st.session_state.get(success_flag):
                st.success(f"✅ '{filename}' uploaded successfully!")
            else:
                st.info("File already uploaded in this session. If you want to re-upload, delete the previous file first.")
        else:
            try:
                data_bytes = uploaded_file.read()
                content = ""

                file_ext = filename.lower().split(".")[-1]

                if file_ext == "pdf":
                    from PyPDF2 import PdfReader
                    reader = PdfReader(io.BytesIO(data_bytes))
                    content = "\n".join([p.extract_text() or "" for p in reader.pages])

                elif file_ext == "docx":
                    from docx import Document
                    doc = Document(io.BytesIO(data_bytes))
                    content = "\n".join([p.text for p in doc.paragraphs])

                else:
                    try:
                        content = data_bytes.decode("utf-8")
                    except Exception:
                        content = data_bytes.decode("latin-1", errors="ignore")

                from utils.file_handle import upload_file, safe_text
                content = safe_text(content)
                if not content.strip():
                    content = "-"

                # Upload and mark session state
                upload_file(selected_bot_id, filename, content)
                st.session_state[flag] = True
                st.session_state[success_flag] = True

                st.success(f"✅ '{filename}' uploaded successfully!")
                st.rerun()

            except Exception as e:
                st.error(f"Upload failed: {type(e).__name__}: {e}")
                if flag in st.session_state:
                    del st.session_state[flag]
    else:
        # Clear flags if no file is currently selected
        for key in list(st.session_state.keys()):
            if key.startswith("uploaded_") or key.startswith("success_"):
                del st.session_state[key]

    # List existing RAG files
    st.subheader("Current RAG Files")
    file_list = selected_bot_info.get("file_data", []) or []

    for idx, file_entry in enumerate(file_list):
        col_name, col_del = st.columns([4, 1])
        col_name.write(file_entry.get("name"))

        delete_key = f"delete_file_{idx}_{selected_bot_id}"
        if col_del.button("🗑️ Delete", key=delete_key):
            # call delete helper to remove from Firestore
            try:
                from utils.file_handle import delete_file
                deleted = delete_file(selected_bot_id, file_entry.get("name"))
            except Exception as e:
                st.error(f"Delete failed: {e}")
                deleted = False

            # clear only that file's session-state upload-flag so user can re-upload
            flag = f"uploaded_{selected_bot_id}_{file_entry.get('name')}"
            # after fetching file_list:
            for fe in file_list:
                flag = f"uploaded_{selected_bot_id}_{fe.get('name')}"
                if st.session_state.get(flag):
                    del st.session_state[flag]
            if flag in st.session_state:
                del st.session_state[flag]

            if deleted:
                st.success("File deleted!")
            else:
                st.warning("File not found / already deleted.")

            # refresh UI (fetches current Firestore state on next run)
            st.rerun()


    # --- Website URLs ---
    st.markdown("---")
    st.subheader("🌐 Manage Website URLs")
    # Add Website URL
    new_url = st.text_input("Add Website URL", key=f"url_{selected_bot_id}")
    if st.button("Add URL", key=f"add_url_{selected_bot_id}"):
        if new_url:
            from utils.file_handle import scrape_and_add_url  # we’ll create this
            scrape_and_add_url(selected_bot_id, new_url)
            st.success(f"URL added and content appended: {new_url}")
            st.rerun()
    
    # Show existing URLs
    # Manage URLs
    st.subheader("Current URLs")
    # Fetch fresh data
    bot_doc = db.collection("bots").document(selected_bot_id).get()
    urls = bot_doc.to_dict().get("config", {}).get("urls", [])
    for u in urls:
        col1, col2 = st.columns([5, 1])
        col1.write(u)
        if col2.button("🗑️ Delete", key=f"del_url_{u}_{selected_bot_id}"):
            from utils.file_handle import delete_url
            delete_url(selected_bot_id, u)
            st.success(f"Deleted {u}")
            st.rerun()


    # --- Embed snippet ---
    st.markdown("---")
    st.write("📄 **Embed this bot on your website:**")
    api_key = selected_bot_info.get("api_key")
    if api_key:
        embed_code = f'<script src="{BACKEND}/static/embed.js" data-api-key="{api_key}" data-bot-name="{selected_bot_info["name"]}"></script>'
        st.code(embed_code, language="html")
        st.markdown(f"""
        **How to use this snippet:**
        1. Copy the code above.
        2. Paste it **before the closing `</body>` tag** in your HTML.
        3. Refresh your website. The chat widget for **{selected_bot_info['name']}** will appear.
        """)
    else:
        st.warning("API key not found for this bot.")

# ---------------- MAIN APP ----------------
def main():

    tabs = st.tabs(["➕ Create Bot", "🛠️ Manage Bots", "💬 My Bots", "👤 Account", "👨‍💻 Meet Us", "📱 WhatsApp Integration"])

    with tabs[0]:
        create_and_save_bot()
    with tabs[1]:
        bot_management_ui()
    with tabs[2]:
        chat_interface()
    with tabs[3]:
        try:
            user
        except NameError:
            user = None
        username = user.get("displayName", user.get("email", "User")) if user else "User"
        st.success(f"👋 Welcome, 👤{username}")
        if st.button("🚪 Sign Out"):
            st.session_state["user"] = None
            cookies.delete("user_email")
            cookies.delete("user_uid")
            cookies.save()
            st.rerun()
    with tabs[4]:
        st.header("👨‍💻 Meet the Mentesa Team")
        st.markdown("---")
        st.markdown("""
        ### 🧠 About Mentesa  
        Mentesa is a **no-code platform** that empowers anyone to create, manage, and chat with their own AI-powered bots.  
        Our mission is to make generative AI accessible, personal, and fun.  
        ### 👥 Our Team  
           - **Mayur Koli** – Founder & Lead Developer  
           - **Anirudh Kapurkar** – Frontend Developer  
           - **Niharika Wagh** – Backend Developer & Research Associate  
        We're constantly innovating to bring smarter, more personalized AI experiences to you.
        """) 
        st.markdown("---")
        st.subheader("🌐 More About Us")    
        st.markdown(
            """
            <a href="https://developer.mentesa.live/" target="_blank">
                <button style="
                    background-color: transparent;
                    color: white;
                    padding: 10px 18px;
                    text-align: center;
                    text-decoration: none;
                    display: inline-block;
                    font-size: 16px;
                    border-radius: 8px;
                    cursor: pointer;
                ">
                    🚀 Team Mentesa
                </button>
            </a>
            """,
            unsafe_allow_html=True
        )
    with tabs[5]:
        # ✅ RESET OTP STATE (forces fresh Step 3 flow)
        for key in ["wa_phone", "verified_phone_id"]:
            if key in st.session_state:
                del st.session_state[key]
        
        st.header("📱 Connect WhatsApp to Your Bot")
        st.markdown("---")

        user = st.session_state.get("user") or {}
        owner_uid = user.get("uid")
        owner_email = user.get("email")

        if not owner_uid or not owner_email:
            st.error("User not found. Please sign in again.")
            st.stop()

        # Load user's bots
        bots = load_user_bots()
        if not bots:
            st.info("You have no bots yet. Create a bot first.")
            st.stop()

        # -------------------------
        # Step 1 — Enter Phone Number ID
        # -------------------------
        st.subheader("Step 1 — Enter Your Phone Number ID")

        phone_number_id = st.text_input(
            "Phone Number ID",
            placeholder="854485194419931"
        )

        method = st.selectbox("OTP Method", ["sms"])

        if st.button("📩 Send OTP"):
            if not phone_number_id.strip():
                st.warning("Enter Phone Number ID first.")
                st.stop()

            try:
                r = requests.post(
                    f"{BACKEND}/wa/register",
                    json={"phone_number_id": phone_number_id.strip(), "method": method}
                )
                if r.status_code == 200:
                    st.success("✅ OTP sent!")
                    st.session_state["wa_phone_id"] = phone_number_id.strip()
                else:
                    st.error(r.text)
            except Exception as e:
                st.error(e)


        # -------------------------
        # Step 2 — Verify OTP
        # -------------------------
        if "wa_phone_id" in st.session_state:
            st.subheader("Step 2 — Verify OTP")
            code = st.text_input("OTP Code", placeholder="123456")

            if st.button("✅ Verify OTP"):
                if not code.strip():
                    st.warning("Enter OTP.")
                    st.stop()

                try:
                    r = requests.post(
                        f"{BACKEND}/wa/verify_otp",
                        json={
                            "phone_number_id": st.session_state["wa_phone_id"],
                            "code": code.strip()
                        }
                    )
                    if r.status_code == 200:
                        st.success("✅ Number verified successfully!")
                        st.session_state["verified_phone_id"] = st.session_state["wa_phone_id"]
                    else:
                        st.error(r.text)
                except Exception as e:
                    st.error(e)


        # -------------------------
        # Step 3 — Select Bot + Connect
        # -------------------------
        if "verified_phone_id" in st.session_state:
            st.subheader("Step 3 — Select Bot to Connect")

            selected_bot = st.selectbox(
                "Choose bot",
                options=bots,
                format_func=lambda b: f"{b.get('name')} ({b.get('id')[:6]})"
            )

            if selected_bot:
                bot_id = selected_bot["id"]

                # Fetch bot API key
                r = requests.get(f"{BACKEND}/bots/{bot_id}/apikey")
                api_key = r.json().get("api_key")

                if st.button("🔗 Connect WhatsApp to Bot"):
                    payload = {
                        "phone_number_id": st.session_state["verified_phone_id"],
                        "api_key": api_key,
                        "bot_id": bot_id,
                        "owner_uid": owner_uid,
                        "owner_email": owner_email,
                    }

                    r = requests.post(f"{BACKEND}/whatsapp/connect", json=payload)

                    if r.status_code == 200:
                        st.success("✅ WhatsApp connected successfully!")
                        st.json(r.json())
                    else:
                        st.error(r.text)


if __name__ == "__main__":
    main()
